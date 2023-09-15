import calendar
import datetime
from io import BytesIO

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import redirect, render
from django.urls import reverse_lazy
from django.views.generic.edit import CreateView
from docx import Document

from .forms import NameForm, NameForm23, NameForm26
from .models import Table, Table23, Table26


@login_required
def home(request):
    return render(request, "home/home.html", {'home': home})

def prehome(request):
    return render(request, "home/prehome.html", {"prehome": prehome})

# Таблица №20
@login_required
def new_form(request):
    
    if request.method == 'POST':
        form = NameForm(request.POST)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.author = request.user
            table = form.save()
            return redirect('/table/' + str(table.id) + '/')
    else:
        form = NameForm()
     

    return render(request, "form/form.html", {'form': form})

@login_required
def table_page(request, pk):
    table = Table.objects.get(id=pk)
    return render(request, "form/table_page.html", {'table': table})

@login_required
def edit_form(request, pk):
    table = Table.objects.get(id=pk)
    
    if request.method == 'POST':
        form = NameForm(request.POST, instance=table)
        if form.is_valid():
            form.author = request.user
            form.save()
            return redirect('/table/' + str(table.id) + '/')
    else:
        form = NameForm(instance=table)
     
    return render(request, "form/edit_form.html", {'form': form})

# Таблица №23
@login_required
def new_form23(request):
    
    if request.method == 'POST':
        form = NameForm23(request.POST)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.subject = request.user
            table = form.save()
            return redirect('/table23/' + str(table.id) + '/')
    else:
        form = NameForm23()
     

    return render(request, "table23/form.html", {'form': form})

@login_required
def table_page23(request, pk):
    table = Table23.objects.get(id=pk)
    return render(request, "table23/table_page23.html", {'table': table})

@login_required
def edit_form23(request, pk):
    table = Table23.objects.get(id=pk)
    
    if request.method == 'POST':
        form = NameForm23(request.POST, instance=table)
        if form.is_valid():
            form.author = request.user
            form.save()
            return redirect('/table23/' + str(table.id) + '/')
    else:
        form = NameForm23(instance=table)
     
    return render(request, "table23/edit_form23.html", {'form': form})


# Таблица №26
@login_required
def new_form26(request):
    
    if request.method == 'POST':
        form = NameForm26(request.POST)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.organisation = request.user
            table = form.save()
            return redirect('/table26/' + str(table.id) + '/')
    else:
        form = NameForm26()
     

    return render(request, "table26/form.html", {'form': form})

@login_required
def table_page26(request, pk):
    table = Table26.objects.get(id=pk)
    return render(request, "table26/table_page26.html", {'table': table})

@login_required
def edit_form26(request, pk):
    table = Table26.objects.get(id=pk)
    
    if request.method == 'POST':
        form = NameForm26(request.POST, instance=table)
        if form.is_valid():
            form.author = request.user
            form.save()
            return redirect('/table26/' + str(table.id) + '/')
    else:
        form = NameForm26(instance=table)
     
    return render(request, "table26/edit_form26.html", {'form': form})

@login_required
def export_data_to_word_for_table20(request):

     # Get your queryset of data 
    queryset = Table.objects.all() 
 
    # Create a new Word document 
    document = Document() 
 
    # Iterate over the data and group it by month 
    data_by_month20 = {} 
    for obj in queryset: 
        month = obj.date.month 
        if month not in data_by_month20: 
            data_by_month20[month] = [] 
        data_by_month20[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month20.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        # Add a table to the document with headers 
        table = document.add_table(rows=2, cols=4)
        hdr_cells = table.rows[0].cells
        table.style = 'Table Grid'
        hdr_cells[0].text = 'Ф.И.О. автора'
        hdr_cells[1].text = '№ патента'
        hdr_cells[2].text = 'Год выдачи'
        hdr_cells[3].text = 'Название'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells = table.add_row().cells 
            row_cells[0].text = str(obj.author_name) 
            row_cells[1].text = str(obj.patent) 
            row_cells[2].text = str(obj.year) 
            row_cells[3].text = str(obj.title) 
 
    # Set the file name and Content-Type header for the response 
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document') 
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx' 
 
    # Save the Word document to the response 
    document.save(response) 
    return response

#------------------------------------------------------------------------------------------
@login_required
def export_data_to_word_for_table23(request):
    # Get your queryset of data 
    queryset23 = Table23.objects.all() 
 
    # Create a new Word document 
    document = Document() 
 
    # Iterate over the data and group it by month 
    data_by_month23 = {} 
    for obj in queryset23: 
        month = obj.date.month 
        if month not in data_by_month23: 
            data_by_month23[month] = [] 
        data_by_month23[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month23.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        # Add a table to the document with headers 
        table23 = document.add_table(rows=2, cols=6)
        hdr_cells23 = table23.rows[0].cells
        table23.style = 'Table Grid'
        hdr_cells23[0].text = 'Предмет международного договора или проекта'
        hdr_cells23[1].text = 'Наименование учебного заведения-партнера'
        hdr_cells23[2].text = 'Дата заключения проектов и/илидоговоров'
        hdr_cells23[3].text = 'дата начала'
        hdr_cells23[4].text = 'дата окончания'
        hdr_cells23[5].text = 'Фактическое наличие на момент проф. контроля'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells23 = table23.add_row().cells
            row_cells23[0].text = str(obj.subject)
            row_cells23[1].text = str(obj.name_of_partner)
            date_str = obj.start_date.strftime("%Y-%m-%d")
            date_str1 = obj.date_of_contract.strftime("%Y-%m-%d")
            date_str2 = obj.end_date.strftime("%Y-%m-%d")
            row_cells23[2].text = date_str1
            row_cells23[3].text = date_str
            row_cells23[4].text = date_str2
            row_cells23[5].text = str(obj.availability) 
 
    # Set the file name and Content-Type header for the response 
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document') 
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx' 
 
    # Save the Word document to the response 
    document.save(response) 
    return response
    


#------------------------------------------------------------------------------------------
@login_required
def export_data_to_word_for_table26(request):
    # Get your queryset of data 
    queryset26 = Table26.objects.all() 
 
    # Create a new Word document 
    document = Document() 
 
    # Iterate over the data and group it by month 
    data_by_month26 = {} 
    for obj in queryset26: 
        month = obj.date.month 
        if month not in data_by_month26: 
            data_by_month26[month] = [] 
        data_by_month26[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month26.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        table26 = document.add_table(rows=2, cols=5)
        hdr_cells26 = table26.rows[0].cells
        table26.style = 'Table Grid'
        hdr_cells26[0].text = 'Организация'
        hdr_cells26[1].text = 'Предмет Соглашения или договора'
        hdr_cells26[2].text = 'На какие направления подготовки/специальности распространяется действие договора'
        hdr_cells26[3].text = 'Дата подписания'
        hdr_cells26[4].text = 'Сроки действия договора'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells26 = table26.add_row().cells
            row_cells26[0].text = str(obj.organisation)
            row_cells26[1].text = str(obj.subject_of_contract)
            row_cells26[2].text = str(obj.directon_of_speciality)
            date_str = obj.date_of_conclusion_of_the_contract.strftime("%Y-%m-%d")
            row_cells26[3].text = str(obj.date_of_conclusion_of_the_contract)
            row_cells26[4].text = str(obj.terms_of_the_contract)
 
    # Set the file name and Content-Type header for the response 
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document') 
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx' 
 
    # Save the Word document to the response 
    document.save(response) 
    return response
    

#------------------------------------------------------------------------------------------

@login_required
def export_data_to_word(request):
    

    # Create a new Word document
    document = Document()



    # Таблица №20

    # Get your queryset of data

    document.add_paragraph('Перечень патентов, полученных в 20____-20___ годы', style='Heading 1')

    queryset = Table.objects.all()

    # Iterate over the data and group it by month 
    data_by_month20 = {} 
    for obj in queryset: 
        month = obj.date.month 
        if month not in data_by_month20: 
            data_by_month20[month] = [] 
        data_by_month20[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month20.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        # Add a table to the document with headers 
        table = document.add_table(rows=2, cols=4)
        hdr_cells = table.rows[0].cells
        table.style = 'Table Grid'
        hdr_cells[0].text = 'Ф.И.О. автора'
        hdr_cells[1].text = '№ патента'
        hdr_cells[2].text = 'Год выдачи'
        hdr_cells[3].text = 'Название'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells = table.add_row().cells 
            row_cells[0].text = str(obj.author_name) 
            row_cells[1].text = str(obj.patent) 
            row_cells[2].text = str(obj.year) 
            row_cells[3].text = str(obj.title) 
 

    # Таблица №23
    document.add_paragraph('Договора о международном сотрудничестве', style='Heading 1')
    # Get your queryset of data
    queryset23 = Table23.objects.all()

        # Iterate over the data and group it by month 
    data_by_month23 = {} 
    for obj in queryset23: 
        month = obj.date.month 
        if month not in data_by_month23: 
            data_by_month23[month] = [] 
        data_by_month23[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month23.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        # Add a table to the document with headers 
        table23 = document.add_table(rows=2, cols=6)
        hdr_cells23 = table23.rows[0].cells
        table23.style = 'Table Grid'
        hdr_cells23[0].text = 'Предмет международного договора или проекта'
        hdr_cells23[1].text = 'Наименование учебного заведения-партнера'
        hdr_cells23[2].text = 'Дата заключения проектов и/илидоговоров'
        hdr_cells23[3].text = 'дата начала'
        hdr_cells23[4].text = 'дата окончания'
        hdr_cells23[5].text = 'Фактическое наличие на момент проф. контроля'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells23 = table23.add_row().cells
            row_cells23[0].text = str(obj.subject)
            row_cells23[1].text = str(obj.name_of_partner)
            date_str = obj.start_date.strftime("%Y-%m-%d")
            date_str1 = obj.date_of_contract.strftime("%Y-%m-%d")
            date_str2 = obj.end_date.strftime("%Y-%m-%d")
            row_cells23[2].text = date_str1
            row_cells23[3].text = date_str
            row_cells23[4].text = date_str2
            row_cells23[5].text = str(obj.availability)

    
    # Таблица №26
    document.add_paragraph('Соглашения о сотрудничестве с организациями образования или научными или научно-образовательными или научно-производственными центрами (магистратура, докторантура)', style='Heading 1')
    # Get your queryset of data
    queryset26 = Table26.objects.all()

   # Iterate over the data and group it by month 
    data_by_month26 = {} 
    for obj in queryset26: 
        month = obj.date.month 
        if month not in data_by_month26: 
            data_by_month26[month] = [] 
        data_by_month26[month].append(obj) 
 
    # Add a table for each month's data 
    for month, data in data_by_month26.items(): 
        # Add a section heading with the month 
        document.add_heading(calendar.month_name[month], level=1) 
 
        table26 = document.add_table(rows=2, cols=5)
        hdr_cells26 = table26.rows[0].cells
        table26.style = 'Table Grid'
        hdr_cells26[0].text = 'Организация'
        hdr_cells26[1].text = 'Предмет Соглашения или договора'
        hdr_cells26[2].text = 'На какие направления подготовки/специальности распространяется действие договора'
        hdr_cells26[3].text = 'Дата подписания'
        hdr_cells26[4].text = 'Сроки действия договора'
 
        # Add data rows to the table 
        for obj in data: 
            row_cells26 = table26.add_row().cells
            row_cells26[0].text = str(obj.organisation)
            row_cells26[1].text = str(obj.subject_of_contract)
            row_cells26[2].text = str(obj.directon_of_speciality)
            date_str = obj.date_of_conclusion_of_the_contract.strftime("%Y-%m-%d")
            row_cells26[3].text = str(obj.date_of_conclusion_of_the_contract)
            row_cells26[4].text = str(obj.terms_of_the_contract)

       

    # Set the file name and Content-Type header for the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx'

    # Save the Word document to the response
    document.save(response)
    return response


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')
        else:
            return render(request, 'login\login.html', {'error_message': 'Invalid login'})
    else:
        return render(request, 'login\login.html')

@login_required
def profile_view(request):
    return render(request, 'login\profile.html', {'user': request.user})

@login_required
def logout_view(request):
    logout(request)
    return redirect('prehome')

class SignUpView(CreateView):
    form_class = UserCreationForm
    success_url = reverse_lazy('login')
    template_name = 'login\signup.html'